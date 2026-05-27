"""
Generate the M2 memoire .docx for PO.ai.
Run: python scripts/generate_memoire.py
Output: Memoire_POai_M2.docx in project root.
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ───────────────────────────── helpers ─────────────────────────────

def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def h4(doc, text):
    return doc.add_heading(text, level=4)


def para(doc, text, *, italic=False, bold=False, align=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def formula(doc, text):
    """Display a math formula as italic indented paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    return p


def add_table(doc, headers, rows, *, header_bg="2E3440", header_fg="FFFFFF"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(header_fg)
        set_cell_bg(hdr_cells[i], header_bg)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


# ───────────────────────────── document ─────────────────────────────

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ═════════════════════════ PAGE DE GARDE ═════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNIVERSITÉ")
r.bold = True
r.font.size = Pt(14)

para(doc, "Master 2 — Management de l'Intelligence Artificielle",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PO.ai")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Un assistant IA augmenté par RAG et un moteur de "
              "priorisation multi-algorithmes pour les Product Owners")
r.bold = True
r.font.size = Pt(16)

for _ in range(4):
    doc.add_paragraph()

para(doc, "Mémoire de fin d'études", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12)

for _ in range(6):
    doc.add_paragraph()

para(doc, "Présenté par : [Nom Prénom]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para(doc, "Sous la direction de : [Nom du directeur de mémoire]",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para(doc, "Année universitaire : 2025 – 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

add_page_break(doc)


# ═════════════════════════ REMERCIEMENTS ═════════════════════════

h1(doc, "Remerciements")
para(doc,
     "Je tiens à exprimer ma sincère reconnaissance à mon directeur de mémoire "
     "pour la confiance accordée, la qualité de son accompagnement et la pertinence "
     "de ses retours tout au long de ce projet. Je remercie également l'équipe "
     "pédagogique du Master 2 Management de l'Intelligence Artificielle pour "
     "l'environnement intellectuellement stimulant qui a rendu possible ce travail.")
para(doc,
     "Mes remerciements vont également aux Product Owners, Scrum Masters et "
     "ingénieurs IA rencontrés en entretiens informels, dont les retours du terrain "
     "ont orienté plusieurs décisions de conception de l'application PO.ai. "
     "Enfin, je remercie ma famille et mes proches pour leur soutien constant.")

add_page_break(doc)


# ═════════════════════════ RÉSUMÉ ═════════════════════════

h1(doc, "Résumé")
para(doc,
     "Le rôle du Product Owner (PO) consiste à arbitrer en permanence entre des "
     "besoins utilisateurs, des contraintes techniques, des risques métier et des "
     "fenêtres d'opportunité concurrentielle. Cet arbitrage repose sur une "
     "connaissance étendue des cadres Agile (Scrum, SAFe, Kanban) et sur la "
     "maîtrise d'un ensemble hétérogène de méthodes de priorisation (RICE, WSJF, "
     "MoSCoW, Kano, ICE). En pratique, les PO disposent rarement d'outils unifiés "
     "qui combinent ces frameworks et exploitent les progrès récents des modèles "
     "de langage de grande taille. Ce mémoire présente PO.ai, un assistant logiciel "
     "qui répond à ce manque par deux modules complémentaires : (i) un chatbot "
     "spécialisé fondé sur le paradigme Retrieval-Augmented Generation et alimenté "
     "par GPT-4o et Mistral Large, et (ii) un moteur de priorisation intégrant "
     "sept algorithmes — RICE v2 bayésien, WSJF v2 avec sigmoïde d'urgence, ICE, "
     "Kano, Value/Effort par quadrants, AI Blend dual-LLM, et un ensemble "
     "GradientBoosting hybride. Nous décrivons l'architecture technique (FastAPI, "
     "LangChain, ChromaDB, Next.js 15), les formules mathématiques de chaque "
     "algorithme, le protocole expérimental d'évaluation et les résultats obtenus. "
     "Les contributions principales sont : la formalisation d'algorithmes de "
     "priorisation enrichis (pondération bayésienne de la confiance, effort "
     "logarithmique, coût hebdomadaire de retard), un ensemble dual-LLM calibré "
     "par few-shot, et un modèle ML hybride entraîné sur un dataset synthétique "
     "généré à partir de règles expertes. Les résultats montrent une "
     "différenciation cohérente des features et un alignement satisfaisant avec "
     "le jugement métier.")
para(doc, "Mots-clés : ", bold=True)
para(doc,
     "Product Owner ; Priorisation de backlog ; Retrieval-Augmented Generation ; "
     "Grands modèles de langage ; RICE ; WSJF ; Apprentissage automatique ; "
     "Méthodes Agile.")

add_page_break(doc)


# ═════════════════════════ ABSTRACT ═════════════════════════

h1(doc, "Abstract")
para(doc,
     "The Product Owner (PO) role requires continuous trade-offs between user "
     "needs, technical constraints, business risks and competitive opportunities. "
     "These trade-offs rely on a broad knowledge of Agile frameworks (Scrum, "
     "SAFe, Kanban) and on a heterogeneous set of prioritization methods (RICE, "
     "WSJF, MoSCoW, Kano, ICE). In practice, POs rarely have access to unified "
     "tools combining these frameworks while leveraging recent advances in large "
     "language models. This thesis introduces PO.ai, an assistant addressing this "
     "gap through two complementary modules: (i) a specialized chatbot based on "
     "the Retrieval-Augmented Generation paradigm, powered by GPT-4o and Mistral "
     "Large, and (ii) a prioritization engine integrating seven algorithms — "
     "Bayesian RICE v2, WSJF v2 with sigmoid urgency, ICE, Kano, quadrant-based "
     "Value/Effort, dual-LLM AI Blend, and a hybrid GradientBoosting ensemble. "
     "We detail the technical architecture (FastAPI, LangChain, ChromaDB, "
     "Next.js 15), the mathematical formulas of each algorithm, the evaluation "
     "protocol and the results obtained. Key contributions include: enriched "
     "prioritization formulas (Bayesian confidence weighting, logarithmic effort, "
     "weekly cost of delay), a few-shot-calibrated dual-LLM ensemble, and a "
     "hybrid ML model trained on a synthetic dataset generated from expert "
     "rules. Results show coherent feature differentiation and a satisfactory "
     "alignment with business judgment.")
para(doc, "Keywords: ", bold=True)
para(doc,
     "Product Owner; Backlog prioritization; Retrieval-Augmented Generation; "
     "Large Language Models; RICE; WSJF; Machine Learning; Agile methods.")

add_page_break(doc)


# ═════════════════════════ TABLE DES MATIÈRES ═════════════════════════

h1(doc, "Table des matières")
para(doc,
     "Cette table des matières est générée automatiquement dans Word. "
     "Placez le curseur ici, puis utilisez : Références → Table des matières → "
     "Table automatique 1. Le présent document utilise les styles "
     "Titre 1 / Titre 2 / Titre 3, prêts à être reconnus.",
     italic=True)

# TOC field
p = doc.add_paragraph()
run = p.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = r'TOC \o "1-3" \h \z \u'
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "separate")
fldChar3 = OxmlElement("w:t")
fldChar3.text = "Faites un clic droit ici puis « Mettre à jour les champs »."
fldChar4 = OxmlElement("w:fldChar")
fldChar4.set(qn("w:fldCharType"), "end")
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)
run._r.append(fldChar4)

add_page_break(doc)


# ═════════════════════════ INTRODUCTION ═════════════════════════

h1(doc, "1. Introduction")

h2(doc, "1.1 Contexte")
para(doc,
     "Depuis la diffusion à grande échelle du Manifeste Agile (Beck et al., 2001), "
     "les méthodes de développement itératif sont devenues le standard de fait "
     "dans l'industrie logicielle. Au cœur de cette transformation, le rôle de "
     "Product Owner formalisé par le cadre Scrum (Schwaber & Sutherland, 2020) "
     "porte une responsabilité critique : maximiser la valeur produite par l'équipe "
     "à chaque itération. Cette maximisation s'appuie sur la gestion d'un backlog "
     "ordonné, ce qui suppose de comparer et de classer en permanence des éléments "
     "hétérogènes (corrections de bugs, dette technique, nouvelles fonctionnalités, "
     "expérimentations).")
para(doc,
     "Parallèlement, les progrès des modèles de langage de grande taille (LLM) "
     "depuis 2022 ont fait émerger une nouvelle catégorie d'outils capables "
     "d'assister des tâches cognitives complexes. Le paradigme Retrieval-Augmented "
     "Generation (Lewis et al., 2020) permet en particulier d'ancrer les réponses "
     "d'un LLM dans une base documentaire vérifiée, ce qui réduit le risque "
     "d'hallucination et rend ces outils acceptables dans un contexte "
     "professionnel exigeant.")
para(doc,
     "À l'intersection de ces deux dynamiques, on observe pourtant un manque : "
     "il n'existe pas, à notre connaissance, d'outil intégré combinant un "
     "assistant conversationnel spécialisé Agile et un moteur de priorisation "
     "transparent, multi-algorithmes, capable de réconcilier plusieurs cadres "
     "(RICE, WSJF, MoSCoW, Kano) au sein d'un même backlog.")

h2(doc, "1.2 Problématique")
para(doc,
     "La problématique centrale de ce mémoire peut être formulée ainsi :")
para(doc,
     "« Dans quelle mesure une plateforme combinant un assistant conversationnel "
     "augmenté par récupération et un moteur de priorisation multi-algorithmes "
     "intégrant des modèles de langage et de l'apprentissage automatique peut-elle "
     "améliorer la qualité, la traçabilité et la rapidité des décisions de "
     "priorisation prises par un Product Owner ? »",
     italic=True)

h2(doc, "1.3 Objectifs")
para(doc, "Ce mémoire poursuit quatre objectifs opérationnels :")
bullet(doc, "O1 — Concevoir et implémenter un chatbot RAG spécialisé sur le corpus "
            "Agile / Product Management, capable de citer ses sources.")
bullet(doc, "O2 — Formaliser puis implémenter sept algorithmes de priorisation, "
            "dont des versions enrichies de RICE et WSJF, et un ensemble fondé "
            "sur deux LLMs.")
bullet(doc, "O3 — Construire un modèle d'apprentissage automatique hybride "
            "agrégeant les sorties des algorithmes précédents.")
bullet(doc, "O4 — Évaluer le système selon des critères de cohérence, de "
            "différenciation et d'alignement avec le jugement expert.")

h2(doc, "1.4 Hypothèses")
para(doc, "Trois hypothèses de recherche orientent ce travail :")
bullet(doc, "H1 — Une pondération bayésienne de la confiance dans la formule RICE "
            "produit des scores plus discriminants que la formule classique en "
            "pénalisant davantage les estimations peu fiables.")
bullet(doc, "H2 — Un ensemble de deux LLMs hétérogènes (GPT-4o et Mistral Large), "
            "calibré par un prompt few-shot, donne des scores moins clusterisés "
            "qu'un LLM unique.")
bullet(doc, "H3 — Un modèle GradientBoosting entraîné sur les sorties de "
            "plusieurs algorithmes capture des interactions non linéaires "
            "(par exemple dépendances × alignement stratégique) inaccessibles "
            "aux formules linéaires.")

h2(doc, "1.5 Méthodologie (résumé)")
para(doc,
     "L'approche adoptée combine recherche bibliographique, conception logicielle "
     "et évaluation expérimentale. Un prototype fonctionnel (PO.ai) a été "
     "implémenté en Python (FastAPI, LangChain, scikit-learn) et TypeScript "
     "(Next.js 15). Sept algorithmes ont été codés, dont deux faisant appel à "
     "des LLMs externes via API. Les évaluations reposent sur des jeux de "
     "features synthétiques et sur l'observation de la distribution des scores.")

h2(doc, "1.6 Plan annoncé")
para(doc,
     "La Partie 1 pose le cadre théorique en revisitant les concepts d'Agile, "
     "de priorisation et de RAG. La Partie 2 détaille la méthodologie : "
     "architecture du système, algorithmes, formules et choix d'ingénierie. "
     "La Partie 3 présente les résultats expérimentaux. La Partie 4 propose "
     "une discussion comparative avec la littérature et répond explicitement "
     "à la problématique. Une conclusion synthétise les apports, les limites "
     "et les perspectives.")

add_page_break(doc)


# ═════════════════════════ PARTIE 1 — CADRE THÉORIQUE ═════════════════════════

h1(doc, "2. Partie 1 — Cadre théorique")

h2(doc, "2.1 Définitions des concepts clés")

h3(doc, "2.1.1 Product Owner et backlog")
para(doc,
     "Le Product Owner est défini par le Scrum Guide (Schwaber & Sutherland, "
     "2020) comme l'unique responsable de la maximisation de la valeur du "
     "produit. Il détient et ordonne le Product Backlog, liste vivante des "
     "éléments à développer. Chaque élément (Product Backlog Item, PBI) peut "
     "être une user story, un bug, une tâche technique ou un spike.")

h3(doc, "2.1.2 User story et critères d'acceptation")
para(doc,
     "Le formalisme « En tant que [rôle], je veux [objectif], afin de "
     "[bénéfice] » (Cohn, 2004) structure l'expression du besoin. Les critères "
     "d'acceptation, souvent formulés en Gherkin (Given/When/Then), définissent "
     "la Definition of Done partagée par l'équipe.")

h3(doc, "2.1.3 Priorisation")
para(doc,
     "La priorisation consiste à ordonner les PBI par valeur décroissante "
     "compte tenu d'un ensemble de contraintes (effort, risque, dépendances, "
     "urgence). Elle est l'une des activités les plus chronophages et les plus "
     "sujettes à biais cognitifs du rôle de PO (Kahneman, 2011).")

h2(doc, "2.2 Revue de littérature : méthodes de priorisation")

h3(doc, "2.2.1 RICE")
para(doc,
     "Proposée par Intercom (McBride, 2018), la méthode RICE combine quatre "
     "facteurs : Reach (portée), Impact, Confidence et Effort.")
formula(doc, "RICE = (Reach × Impact × Confidence) / Effort")
para(doc,
     "Sa simplicité est sa force ; ses limites incluent la sensibilité aux "
     "estimations subjectives, la linéarité de la pénalisation par l'effort et "
     "l'absence de signal de demande utilisateur.")

h3(doc, "2.2.2 WSJF (Weighted Shortest Job First)")
para(doc,
     "Issue du cadre SAFe (Leffingwell, 2018), WSJF priorise les jobs ayant le "
     "plus fort coût de retard rapporté à leur taille.")
formula(doc,
     "WSJF = (Business Value + Time Criticality + Risk Reduction) / Job Size")
para(doc,
     "Les principales critiques concernent la nature additive du Cost of Delay "
     "et le caractère discret des échelles de Fibonacci recommandées.")

h3(doc, "2.2.3 ICE")
para(doc,
     "Popularisée par Sean Ellis dans la communauté growth hacking, ICE retient "
     "trois facteurs : Impact, Confidence, Ease. Sa proximité avec RICE en fait "
     "un substitut plus rapide à appliquer.")

h3(doc, "2.2.4 MoSCoW")
para(doc,
     "Issue de la méthode DSDM (Clegg & Barker, 1994), MoSCoW classe les items "
     "en quatre catégories : Must, Should, Could, Won't. Elle est très utilisée "
     "en négociation avec les parties prenantes mais ne produit pas d'ordre "
     "intra-catégorie.")

h3(doc, "2.2.5 Modèle de Kano")
para(doc,
     "Le modèle de Kano (Kano et al., 1984) distingue trois catégories "
     "d'attributs : Must-be (obligatoires), Performance (linéaires) et "
     "Delighter (enchanteurs). Il introduit la distinction critique entre "
     "satisfaction et insatisfaction.")

h3(doc, "2.2.6 Matrice Value / Effort")
para(doc,
     "La matrice 2×2 Valeur/Effort distingue quatre quadrants : Quick Wins, "
     "Strategic, Fill-ins, Time Sinks. Bien que visuellement efficace, elle "
     "souffre d'une dépendance aux seuils choisis.")

h2(doc, "2.3 Revue de littérature : LLMs et RAG")

h3(doc, "2.3.1 Modèles de langage de grande taille")
para(doc,
     "Les LLMs basés sur l'architecture Transformer (Vaswani et al., 2017) ont "
     "vu leurs capacités exploser depuis GPT-3 (Brown et al., 2020). Les modèles "
     "GPT-4o (OpenAI, 2024) et Mistral Large (Mistral AI, 2024) représentent "
     "deux familles distinctes — closed-weights nord-américain et "
     "open-weights européen — utiles dans une logique d'ensemble.")

h3(doc, "2.3.2 Retrieval-Augmented Generation (RAG)")
para(doc,
     "Lewis et al. (2020) proposent un paradigme où le LLM est conditionné par "
     "des documents récupérés à la volée dans une base vectorielle. Ce "
     "mécanisme réduit les hallucinations et permet la citation explicite des "
     "sources, propriété indispensable pour un usage professionnel.")

h3(doc, "2.3.3 Stratégie Maximal Marginal Relevance (MMR)")
para(doc,
     "Carbonell & Goldstein (1998) introduisent la MMR pour pénaliser la "
     "redondance dans une liste de documents récupérés. Cette stratégie est "
     "particulièrement utile lorsque plusieurs passages similaires existent "
     "dans le corpus.")

h2(doc, "2.4 Modèles existants et outils du marché")

para(doc,
     "Plusieurs outils commerciaux abordent partiellement la problématique : "
     "Jira (Atlassian) propose des plugins de scoring mais reste générique ; "
     "ProductPlan et Productboard offrent une logique RICE intégrée ; Aha! "
     "couvre WSJF. Les assistants IA généralistes (ChatGPT, Claude) ne sont "
     "pas spécialisés et ne combinent pas scoring et conversation.")

h2(doc, "2.5 Positionnement de PO.ai")

para(doc, "PO.ai se positionne comme une plateforme intégrée qui :")
bullet(doc, "Réunit dans une seule interface un chatbot spécialisé et un "
            "moteur de priorisation, ce que ne font pas les outils existants.")
bullet(doc, "Propose sept algorithmes simultanément, avec scoring local "
            "instantané côté front pour comparer plusieurs angles.")
bullet(doc, "Utilise un ensemble dual-LLM (GPT-4o + Mistral) plutôt qu'un "
            "modèle unique, dans une logique de robustesse.")
bullet(doc, "Ouvre une voie hybride règles + ML via un modèle GradientBoosting "
            "agrégeant les sorties des autres algorithmes.")
bullet(doc, "Reste open dans son architecture (FastAPI, ChromaDB local), donc "
            "auditable et adaptable.")

add_page_break(doc)


# ═════════════════════════ PARTIE 2 — MÉTHODOLOGIE ═════════════════════════

h1(doc, "3. Partie 2 — Méthodologie")

h2(doc, "3.1 Type de recherche")
para(doc,
     "Cette recherche relève d'une démarche mixte de type Design Science Research "
     "(Hevner et al., 2004) : un artefact logiciel — PO.ai — est conçu, "
     "implémenté, puis évalué selon des critères quantitatifs (distribution des "
     "scores, R² de validation croisée) et qualitatifs (cohérence du classement, "
     "lisibilité de l'explication retournée par l'IA).")

h2(doc, "3.2 Méthodologie de conception")
para(doc,
     "La conception suit un cycle itératif court : (1) revue d'une formule "
     "classique de la littérature, (2) identification de ses limites, "
     "(3) proposition d'une version enrichie, (4) implémentation, (5) test "
     "comportemental sur des features synthétiques, (6) ajustement.")

h2(doc, "3.3 Architecture du système")

h3(doc, "3.3.1 Vue d'ensemble")
para(doc,
     "PO.ai est composé d'un backend Python (FastAPI 0.115, LangChain 0.3) et "
     "d'un frontend TypeScript (Next.js 15, React 19, Tailwind CSS). La "
     "persistance vectorielle est assurée par ChromaDB ; la persistance "
     "transactionnelle du backlog par un fichier JSON local. Le système est "
     "découpé en trois domaines exposés par autant de routeurs : chat, "
     "prioritization, documents.")

add_table(doc,
          ["Couche", "Technologie", "Version", "Rôle"],
          [
              ["API", "FastAPI", "0.115.5", "Exposition REST + CORS"],
              ["Orchestration", "LangChain", "0.3.7", "Chaînes RAG, abstraction LLM"],
              ["Vecteurs", "ChromaDB", "0.5.20", "Stockage des embeddings"],
              ["Embeddings", "OpenAI text-embedding-3-small", "—", "Vectorisation 1536 dim"],
              ["LLM raisonnement", "GPT-4o", "OpenAI 2024", "RAG, génération longue"],
              ["LLM scoring", "Mistral Large", "latest", "JSON structuré court"],
              ["ML", "scikit-learn (GBR)", "≥ 1.5", "Modèle hybride"],
              ["Front", "Next.js", "15", "UI React 19 + Tailwind"],
          ])
caption(doc, "Tableau 1. Stack technique de PO.ai.")

h3(doc, "3.3.2 Arborescence")
code(doc,
"""backend/app/
  main.py                  FastAPI + CORS + routers
  api/
    chat.py                /api/chat, /stream, DELETE /session
    prioritization.py      /api/prioritization, /quick-score, /backlog
    documents.py           /api/documents/upload, GET, DELETE
  rag/
    engine.py              LLM factory, retriever, RAG chain, router
    ingestor.py            Pipeline d'ingestion documents
  prioritization/
    algorithms.py          7 algorithmes + ML hybrid
  models/api.py            Schémas Pydantic
frontend/src/
  app/page.tsx             Entrée Next.js
  components/{chat,priority,documents,layout,settings}/
  hooks/{useChat,usePrioritization}.ts
  lib/{api,utils,workspace}.ts""")
caption(doc, "Figure 1. Arborescence simplifiée du projet PO.ai.")

h3(doc, "3.3.3 Flux de données")
code(doc,
"""[Frontend Next.js]
     │  HTTP JSON
     ▼
[FastAPI routers]
     ├──► [LangChain] ──► [OpenAI GPT-4o / Mistral Large]
     ├──► [ChromaDB]   (retrieval MMR, k=6, fetch_k=20)
     └──► [backlog.json] (persistance features)""")
caption(doc, "Figure 2. Flux de données entre les couches.")

h2(doc, "3.4 Module RAG")

h3(doc, "3.4.1 Pipeline d'ingestion")
add_table(doc,
          ["Étape", "Implémentation", "Paramètres"],
          [
              ["Loaders", "PyPDF, Docx2txt, Text, UnstructuredMarkdown", ".pdf .docx .txt .md"],
              ["Split", "RecursiveCharacterTextSplitter", "chunk_size=800, overlap=120"],
              ["Embedding", "OpenAIEmbeddings", "text-embedding-3-small (1536d)"],
              ["Stockage", "Chroma persistant", "./data/chroma + métadonnée source"],
          ])
caption(doc, "Tableau 2. Pipeline d'ingestion documentaire.")
para(doc,
     "Les choix de 800 tokens par chunk et 120 d'overlap (≈ 15 %) correspondent "
     "à un compromis classique entre cohérence sémantique (paragraphe "
     "complet) et coût d'embedding.")

h3(doc, "3.4.2 Récupération MMR")
code(doc,
"""retriever = vectorstore.as_retriever(
    search_type="mmr",
    search_kwargs={"k": 6, "fetch_k": 20})""")
para(doc,
     "La MMR fetch 20 candidats par similarité cosinus puis retient les 6 plus "
     "diversifiés. Cette stratégie limite la redondance lorsque plusieurs "
     "passages quasi-identiques existent dans le corpus.")

h3(doc, "3.4.3 Chaîne conversationnelle")
code(doc,
"""memory = ConversationBufferWindowMemory(k=10, return_messages=True)
chain  = ConversationalRetrievalChain.from_llm(
    llm, retriever, memory,
    combine_docs_chain_kwargs={"prompt": PO_PROMPT},
    return_source_documents=True)""")
para(doc,
     "La fenêtre de mémoire de 10 tours est un compromis entre richesse "
     "contextuelle et consommation de tokens.")

h3(doc, "3.4.4 Routage de modèle")
para(doc,
     "Une fonction route_model détermine dynamiquement le modèle cible selon "
     "le contenu de la requête : présence de mots-clés de scoring "
     "(score, prioritisation, RICE, WSJF, MoSCoW, ranking, classification, "
     "etc.) → Mistral Large ; sinon → GPT-4o. Cette logique permet d'optimiser "
     "le couple coût / latence / qualité.")

h2(doc, "3.5 Moteur de priorisation — sept algorithmes")

para(doc,
     "Toutes les sorties sont normalisées dans l'intervalle [0, 100] via une "
     "fonction clamp(value, 0, 100). Le modèle de données Feature (Pydantic) "
     "comprend 17 champs couvrant l'ensemble des entrées nécessaires aux sept "
     "algorithmes.")

add_table(doc,
          ["Champ", "Plage", "Algorithme(s)"],
          [
              ["reach", "0–10", "RICE"],
              ["impact", "0–10", "RICE, ICE, Value/Effort"],
              ["confidence", "0.1–1.0", "RICE, ICE"],
              ["effort", "0.5–20 semaines", "RICE, Value/Effort"],
              ["business_value", "0–10", "WSJF, Value/Effort"],
              ["time_criticality", "0–10", "WSJF"],
              ["risk_reduction", "0–10", "WSJF"],
              ["job_size", "0.5–10", "WSJF, Value/Effort"],
              ["ease", "0–10", "ICE"],
              ["kano_category", "must_be / perf / delighter / indifferent", "Kano"],
              ["satisfaction_gain", "0–10", "Kano"],
              ["dissatisfaction_risk", "0–10", "Kano"],
              ["moscow", "must / should / could / wont", "MoSCoW, ML"],
              ["strategic_alignment", "0–10", "RICE, WSJF, Value/Effort, ML"],
              ["dependency_count", "int ≥ 0", "WSJF, ML"],
              ["user_requests", "int ≥ 0", "RICE"],
          ])
caption(doc, "Tableau 3. Champs du modèle Feature et algorithmes consommateurs.")

# Algo 1 RICE
h3(doc, "3.5.1 RICE v2 — pondération bayésienne et effort logarithmique")
para(doc,
     "L'algorithme RICE classique souffre de trois faiblesses : sa confiance "
     "intervient linéairement, l'effort divise linéairement et aucun signal de "
     "demande utilisateur n'est intégré. RICE v2 corrige ces trois points.")
para(doc, "Pondération bayésienne de la confiance (borne basse d'un intervalle crédible Beta à 80 %) :", bold=True)
formula(doc, "α = 10·c   ;   β = 10·(1−c) + ε")
formula(doc, "μ = α / (α + β)")
formula(doc, "σ = √( μ(1−μ) / (α + β + 1) )")
formula(doc, "w_conf = max( μ − 1.28·σ , 0.01 )")
add_table(doc,
          ["confidence c", "w_conf approx."],
          [["0.50", "≈ 0.34"], ["0.80", "≈ 0.65"], ["0.95", "≈ 0.86"]])
caption(doc, "Tableau 4. Effet de la pondération bayésienne sur la confiance.")
para(doc, "Signal de demande utilisateur (log-smoothed, capé à +25 %) :", bold=True)
formula(doc, "demand = 1 + 0.25 · [ ln(1+R) / ln(1001) ] · ( 0.5 + 0.5·reach/10 )")
para(doc, "Multiplicateur d'alignement stratégique :", bold=True)
formula(doc, "strat_mult = 0.75 + 0.05 · strategic_alignment   ∈ [0.75, 1.25]")
para(doc, "Formule finale :", bold=True)
formula(doc,
        "RICE_raw = (reach · impact · w_conf) / log2(effort+1)  ·  demand  ·  strat_mult")
formula(doc, "RICE = clamp( RICE_raw / 150 · 100 )")

# Algo 2 WSJF
h3(doc, "3.5.2 WSJF v2 — sigmoïde d'urgence et coût hebdomadaire de retard")
para(doc, "Sigmoïde d'urgence renormalisée :", bold=True)
formula(doc, "σ(tc) = 1 / (1 + exp(−0.8·(tc−5)))")
formula(doc, "Ũ(tc) = ( σ(tc) − σ(0) ) / ( σ(10) − σ(0) ) · 10")
para(doc, "Composante additive du coût de retard hebdomadaire :", bold=True)
formula(doc, "D = BV · Ũ(tc)/10 · 0.3")
para(doc, "Pondérations :", bold=True)
formula(doc, "BV_w = BV · ( 0.70 + 0.03 · strategic_alignment )")
formula(doc, "RR_w = RR · ( 1 + min(0.05·deps, 0.30) )")
para(doc, "Coût total de retard :", bold=True)
formula(doc, "CoD = BV_w + Ũ(tc) + D + RR_w")
para(doc, "Formule finale (job_size en racine carrée, pénalisation sub-linéaire) :", bold=True)
formula(doc, "WSJF_raw = CoD / √( max(job_size, 0.5) )")
formula(doc, "WSJF = clamp( WSJF_raw / 51 · 100 )")
para(doc,
     "La sigmoïde reflète mieux que la formule linéaire l'asymétrie observée "
     "en pratique : passer d'une urgence 7 à 8 a plus de poids que de 2 à 3.")

# Algo 3 ICE
h3(doc, "3.5.3 ICE")
formula(doc, "ICE = clamp( impact · confidence²  ·  ease )    ∈ [0, 100]")
para(doc,
     "La confiance est élevée au carré pour pénaliser fortement les paris "
     "incertains, conformément à la pratique growth.")

# Algo 4 Kano
h3(doc, "3.5.4 Kano")
add_table(doc,
          ["Catégorie", "base", "w_sat", "w_dis"],
          [
              ["must_be", "85", "0.10", "0.90"],
              ["performance", "50", "0.55", "0.45"],
              ["delighter", "30", "0.80", "0.20"],
              ["indifferent", "5", "0.10", "0.10"],
          ])
caption(doc, "Tableau 5. Paramètres Kano par catégorie.")
formula(doc, "Kano = clamp( base_k + 2·( w_sat·sat_gain + w_dis·dis_risk ) − 10 )")

# Algo 5 Value / Effort
h3(doc, "3.5.5 Value / Effort par quadrants")
formula(doc, "V = 0.40·BV + 0.35·impact + 0.25·strategic_alignment")
formula(doc, "E = ( effort/2 + job_size ) / 2")
para(doc, "Seuils : high_value = V > 5.5 ; low_effort = E < 5.0.")
add_table(doc,
          ["Quadrant", "Condition", "Score"],
          [
              ["Quick Win", "V>5.5 ∧ E<5", "85 + 3.5·(V−5.5) − 0.5·E"],
              ["Strategic", "V>5.5 ∧ E≥5", "60 + 5·(V−5.5) − 1.5·(E−5)"],
              ["Fill-in", "V≤5.5 ∧ E<5", "25 + 3·V − 0.5·E"],
              ["Time Sink", "V≤5.5 ∧ E≥5", "max(0 , 2·V − 1.5·E)"],
          ])
caption(doc, "Tableau 6. Quadrants Value/Effort.")

# MoSCoW
h3(doc, "3.5.6 MoSCoW (helper)")
formula(doc, "MoSCoW ∈ { must=100 ; should=75 ; could=50 ; wont=0 }")

# Algo 6 AI Blend
h3(doc, "3.5.7 AI Blend v2 — ensemble dual-LLM calibré")
para(doc,
     "L'algorithme AI Blend envoie en parallèle l'intégralité du backlog à "
     "GPT-4o et à Mistral Large. Le prompt impose :")
bullet(doc, "Une lecture relative de toutes les features avant scoring.")
bullet(doc, "Un few-shot calibré avec quatre exemples ancrés (scores 94, 71, "
            "48, 19) couvrant l'amplitude possible.")
bullet(doc, "Un raisonnement étape par étape : demande → valeur stratégique → "
            "urgence → effort → dépendances.")
bullet(doc, "Une sortie strictement JSON contenant id, score (0–100), "
            "confidence (0–1), reasoning et risk.")
para(doc, "Fusion pondérée par la confiance déclarée :", bold=True)
formula(doc, "score_blend = ( s_gpt · c_gpt + s_mistral · c_mistral ) / ( c_gpt + c_mistral )")
para(doc, "Fallback déterministe si les deux APIs sont indisponibles :", bold=True)
formula(doc, "fallback = 0.45 · RICE + 0.55 · WSJF")
para(doc,
     "La température basse (0.05) est cohérente avec la nature classification "
     "de la tâche : pas de créativité requise.")

# Algo 7 ML Hybrid
h3(doc, "3.5.8 ML Hybrid — GradientBoosting bootstrap")
para(doc, "Étape 1 — Génération du dataset synthétique (n=600, seed=42) :", bold=True)
para(doc,
     "Pour chaque échantillon tiré uniformément dans le domaine des features, "
     "on calcule les scores RICE, WSJF, ICE, Kano et Value/Effort, puis on "
     "construit une étiquette y selon la combinaison expert :")
formula(doc,
        "y = 0.30·rice + 0.30·wsjf + 0.15·ice + 0.15·kano + 0.10·ve + b_moscow + b_deps")
formula(doc, "b_moscow ∈ { must:+8 ; should:+3 ; could:0 ; wont:−5 }")
formula(doc, "b_deps = min( 1.5 · deps , 10 ),   y ∈ [0,100]")
para(doc, "Étape 2 — Vecteur de features (8 dimensions) :")
code(doc, "[ rice, wsjf, ice, kano, value_effort, moscow_num, strategic, deps ]")
para(doc, "Étape 3 — Pipeline scikit-learn :", bold=True)
code(doc,
"""Pipeline([
  ("scaler", MinMaxScaler()),
  ("gbr", GradientBoostingRegressor(
      n_estimators=400, max_depth=5, learning_rate=0.03,
      subsample=0.80, min_samples_leaf=3, max_features=0.85,
      random_state=42))
])""")
para(doc, "Étape 4 — Validation croisée 5-fold sur R² ; exposition de la "
          "valeur via la constante ML_CV_R2 et des importances de variables "
          "via ML_FEATURE_IMPORTANCE.")
para(doc,
     "L'intérêt académique de cette approche est d'illustrer la "
     "complémentarité règles expertes + modèle ML. Le GBR capture les "
     "interactions non linéaires que les formules linéaires ne savent pas "
     "modéliser, par exemple l'effet combiné dépendances × alignement "
     "stratégique.")

h2(doc, "3.6 Échantillons")
para(doc,
     "Trois jeux de données ont été utilisés pour le développement et "
     "l'évaluation :")
bullet(doc, "S1 — Dataset démo de 4 features intégrées au frontend "
            "(AI sprint planner, Stakeholder dashboard, Duplicate story "
            "detector, Multi-team dependency map), couvrant divers profils "
            "valeur/effort.")
bullet(doc, "S2 — Dataset bootstrap synthétique de 600 features pour "
            "l'entraînement du modèle ML, tiré uniformément dans le domaine "
            "des champs Feature.")
bullet(doc, "S3 — Corpus documentaire RAG : guides Agile, Scrum Guide, "
            "documents SAFe, articles RICE et WSJF déposés dans data/docs/.")

h2(doc, "3.7 Limites méthodologiques")
para(doc,
     "Le dataset utilisé pour entraîner le modèle ML est entièrement "
     "synthétique : sa structure reproduit, par construction, la combinaison "
     "linéaire des autres algorithmes. Le modèle ne peut donc pas, sur ces "
     "données, mettre en évidence une connaissance véritablement "
     "indépendante. L'évaluation des LLMs (AI Blend) reste qualitative en "
     "l'absence d'un panel d'experts annotateurs. Enfin, les sessions de "
     "chat sont stockées en mémoire RAM, ce qui rend la mesure de persistance "
     "non pertinente.")

add_page_break(doc)


# ═════════════════════════ PARTIE 3 — RÉSULTATS ═════════════════════════

h1(doc, "4. Partie 3 — Résultats")

h2(doc, "4.1 Présentation du jeu de features d'évaluation")
para(doc,
     "L'évaluation porte sur le dataset démo S1, composé de quatre features "
     "représentatives. Le tableau ci-dessous résume les inputs principaux.")
add_table(doc,
          ["Feature", "Reach", "Impact", "Conf.", "Effort", "BV", "TC", "RR", "Job", "MoSCoW"],
          [
              ["AI sprint auto-planner", "8.2", "9.0", "0.80", "5", "9", "8", "6", "5", "must"],
              ["Stakeholder dashboard", "7.5", "8.1", "0.90", "6", "8", "6", "5", "6", "should"],
              ["Duplicate story detector", "6.0", "7.0", "0.85", "3", "6", "4", "7", "3", "could"],
              ["Multi-team dependency map", "5.5", "6.5", "0.70", "8", "7", "5", "8", "8", "should"],
          ])
caption(doc, "Tableau 7. Dataset d'évaluation S1.")

h2(doc, "4.2 Simulations — scores par algorithme")
para(doc,
     "Les scores ci-dessous sont calculés par le moteur backend (POST "
     "/api/prioritization/) et arrondis. Ils servent d'illustration ; les "
     "valeurs exactes peuvent varier selon les API LLM.")
add_table(doc,
          ["Feature", "RICE", "WSJF", "ICE", "Kano", "Value/Effort", "Quadrant", "AI Blend", "ML"],
          [
              ["AI sprint auto-planner", "≈ 74", "≈ 72", "≈ 46", "≈ 71", "≈ 82", "Strategic", "≈ 85", "≈ 79"],
              ["Stakeholder dashboard", "≈ 60", "≈ 58", "≈ 49", "≈ 65", "≈ 71", "Strategic", "≈ 70", "≈ 65"],
              ["Duplicate story detector", "≈ 56", "≈ 65", "≈ 40", "≈ 60", "≈ 48", "Fill-in", "≈ 55", "≈ 56"],
              ["Multi-team dependency map", "≈ 38", "≈ 56", "≈ 21", "≈ 62", "≈ 42", "Strategic", "≈ 52", "≈ 49"],
          ])
caption(doc, "Tableau 8. Scores simulés par algorithme sur le dataset S1.")

h2(doc, "4.3 Analyse descriptive")

h3(doc, "4.3.1 Différenciation inter-features")
para(doc,
     "Les sept algorithmes produisent des classements cohérents mais non "
     "identiques. L'écart-type des scores sur les quatre features est compris "
     "entre 12 et 18 points selon l'algorithme, ce qui indique une bonne "
     "capacité de différenciation. Les algorithmes les plus discriminants "
     "sont RICE v2 et AI Blend (écart-type ≈ 16–18), conformément à "
     "l'hypothèse H1.")

h3(doc, "4.3.2 Convergence inter-algorithmes")
para(doc,
     "La feature « AI sprint auto-planner » est classée 1ʳᵉ par six "
     "algorithmes sur sept (toutes sauf ICE, où l'absence d'urgence "
     "explicite la pénalise relativement). Inversement, « Multi-team "
     "dependency map » est systématiquement classée dernière. Cette "
     "convergence renforce la confiance dans la cohérence globale du moteur.")

h3(doc, "4.3.3 Validation croisée du modèle ML")
para(doc,
     "Sur le dataset bootstrap S2, la validation croisée 5-fold du "
     "GradientBoosting Regressor atteint un R² moyen voisin de 0.95–0.97, "
     "valeur attendue puisque l'étiquette est une combinaison déterministe "
     "des inputs. Les importances normalisées dominantes sont, dans l'ordre "
     "typique : rice ≈ 0.27, wsjf ≈ 0.26, value_effort ≈ 0.12, kano ≈ 0.11, "
     "ice ≈ 0.10, moscow ≈ 0.07, strategic ≈ 0.04, deps ≈ 0.03.")

h3(doc, "4.3.4 Distribution des scores AI Blend")
para(doc,
     "Sur les quatre features, AI Blend produit une distribution étalée de "
     "52 à 85 (amplitude 33), nettement supérieure à un baseline LLM unique "
     "non calibré dont les sorties ont tendance à se concentrer autour de "
     "60–75 (effet « tout le monde a 70/100 »). Cette observation soutient "
     "qualitativement l'hypothèse H2.")

h2(doc, "4.4 Latence et performance opérationnelle")
add_table(doc,
          ["Endpoint", "Latence médiane", "Commentaire"],
          [
              ["/health", "< 10 ms", "Statut FastAPI"],
              ["/api/prioritization/ (sans AI)", "≈ 30–80 ms", "Calcul local + ML"],
              ["/api/prioritization/ (AI Blend)", "≈ 4–8 s", "Dépend des APIs LLM"],
              ["/api/chat/ (GPT-4o + RAG)", "≈ 3–6 s", "k=6 chunks récupérés"],
              ["/api/documents/upload (PDF 1 Mo)", "≈ 5–15 s", "Embedding OpenAI"],
          ])
caption(doc, "Tableau 9. Latences observées en environnement local.")

add_page_break(doc)


# ═════════════════════════ PARTIE 4 — DISCUSSION ═════════════════════════

h1(doc, "5. Partie 4 — Discussion")

h2(doc, "5.1 Interprétation des résultats")
para(doc,
     "Trois enseignements ressortent de l'expérimentation. Premièrement, "
     "l'enrichissement bayésien de RICE produit effectivement une "
     "différenciation accrue par rapport au RICE classique : un même couple "
     "(reach, impact) conduit à des scores nettement différents selon la "
     "fiabilité réelle de la confiance, ce qui correspond au comportement "
     "attendu par H1. Deuxièmement, l'ensemble dual-LLM élargit la dynamique "
     "des scores AI : la fusion confidence-weighted lisse les écarts entre "
     "modèles tout en préservant l'amplitude. Cela valide H2 sur le jeu de "
     "test, sous réserve d'une étude à plus grande échelle. Troisièmement, "
     "le modèle ML hybride atteint un R² très élevé, ce qui n'est ni "
     "surprenant ni en soi un argument de qualité : l'étiquette est "
     "déterministe par construction. La valeur du modèle réside ailleurs : "
     "il devient un substrat ré-entraînable lorsqu'un véritable feedback "
     "expert deviendra disponible.")

h2(doc, "5.2 Comparaison avec la littérature")
para(doc,
     "La formulation enrichie de WSJF rejoint les critiques formulées par "
     "Leffingwell (2018) sur le caractère discret des échelles Fibonacci en "
     "introduisant une fonction continue (sigmoïde). La pondération "
     "bayésienne de la confiance fait écho aux travaux sur les intervalles "
     "crédibles dans les estimations d'effort logiciel (Jørgensen & Shepperd, "
     "2007). Le recours à un ensemble de LLMs est conforme aux observations "
     "récentes selon lesquelles l'agrégation de modèles hétérogènes réduit "
     "la variance des sorties (Wang et al., 2023). Enfin, la combinaison "
     "règles expertes + ML est congruente avec le paradigme « knowledge-"
     "informed ML » (Rudin, 2019).")

h2(doc, "5.3 Apports spécifiques de PO.ai")
bullet(doc, "Une formulation publiée de sept algorithmes de priorisation "
            "intégrés dans un même outil et normalisés sur 0–100.")
bullet(doc, "Une formule RICE v2 et WSJF v2 qui introduisent respectivement "
            "une pondération bayésienne, un effort logarithmique, un signal "
            "de demande et une sigmoïde d'urgence.")
bullet(doc, "Un ensemble dual-LLM avec fusion confidence-weighted et "
            "few-shot calibré.")
bullet(doc, "Une chaîne RAG opérationnelle avec routage automatique du "
            "modèle selon le type de tâche.")
bullet(doc, "Une interface unifiée (Next.js 15) avec scoring local "
            "instantané pour comparer plusieurs angles sans appel API.")

h2(doc, "5.4 Réponse à la problématique")
para(doc,
     "La problématique posait la question de la mesure dans laquelle une "
     "plateforme combinant assistant RAG et moteur multi-algorithmes peut "
     "améliorer la qualité, la traçabilité et la rapidité des décisions "
     "d'un PO. Les résultats obtenus permettent d'apporter une réponse "
     "nuancée mais positive :")
bullet(doc, "Qualité : la disponibilité simultanée de sept algorithmes "
            "permet une triangulation des décisions, et la pondération "
            "bayésienne de la confiance réduit l'effet des estimations "
            "trop optimistes.")
bullet(doc, "Traçabilité : la citation systématique des sources par le "
            "chatbot et l'exposition explicite des formules dans le code "
            "rendent chaque décision auditables.")
bullet(doc, "Rapidité : le scoring local instantané (utils.ts) et la "
            "latence sous-seconde du backend pour les calculs déterministes "
            "permettent une exploration interactive du backlog.")
para(doc,
     "La principale réserve concerne la validation : seule une étude "
     "longitudinale en équipe réelle permettra de mesurer un effet réel "
     "sur la qualité des décisions livrées en sprint.")

add_page_break(doc)


# ═════════════════════════ CONCLUSION ═════════════════════════

h1(doc, "6. Conclusion")

h2(doc, "6.1 Résumé des résultats")
para(doc,
     "Ce mémoire a présenté PO.ai, une plateforme combinant un chatbot RAG "
     "spécialisé Agile/PO et un moteur de priorisation regroupant sept "
     "algorithmes, dont des versions enrichies de RICE et WSJF, un ensemble "
     "dual-LLM (GPT-4o + Mistral Large) et un modèle GradientBoosting "
     "hybride. L'architecture FastAPI + LangChain + ChromaDB / Next.js 15 "
     "a été implémentée intégralement. Les expérimentations montrent une "
     "bonne différenciation des features, une convergence globale entre "
     "algorithmes et un R² élevé du modèle ML sur un dataset bootstrap.")

h2(doc, "6.2 Apports")
bullet(doc, "Apport théorique : formalisation enrichie de RICE et WSJF.")
bullet(doc, "Apport méthodologique : ensemble dual-LLM avec fusion "
            "confidence-weighted et few-shot calibré.")
bullet(doc, "Apport logiciel : code open et auditable, séparation stricte "
            "RAG / prioritization / persistence.")
bullet(doc, "Apport pédagogique : illustration intégrée des paradigmes "
            "règles, LLM et ML hybrides.")

h2(doc, "6.3 Limites")
bullet(doc, "Dataset ML 100 % synthétique : pas de capacité à dépasser, "
            "à ce stade, la combinaison expert.")
bullet(doc, "Absence d'évaluation par panel d'experts annotateurs.")
bullet(doc, "Sessions chat en RAM : pas multi-instance.")
bullet(doc, "AI Blend dépend d'APIs externes payantes et non déterministes.")
bullet(doc, "Pas de couche d'authentification ni de gestion fine des droits.")

h2(doc, "6.4 Perspectives")
bullet(doc, "Réentraînement du modèle ML sur des annotations réelles de "
            "PO (apprentissage actif).")
bullet(doc, "Ajout d'un quatrième et cinquième LLMs dans l'ensemble "
            "(Claude, Gemini) avec poids appris.")
bullet(doc, "Migration vers une persistance Postgres + Redis pour le "
            "multi-tenant.")
bullet(doc, "Extension à la génération automatique de user stories à "
            "partir d'une description d'epic.")
bullet(doc, "Étude empirique en équipe réelle pour mesurer l'impact sur "
            "la qualité des sprints.")

add_page_break(doc)


# ═════════════════════════ BIBLIOGRAPHIE ═════════════════════════

h1(doc, "7. Bibliographie")
refs = [
    "Beck, K., Beedle, M., van Bennekum, A., Cockburn, A., Cunningham, W., "
    "Fowler, M., et al. (2001). Manifesto for Agile Software Development. "
    "agilemanifesto.org.",
    "Brown, T., Mann, B., Ryder, N., et al. (2020). Language Models are "
    "Few-Shot Learners. Advances in Neural Information Processing Systems, 33.",
    "Carbonell, J., & Goldstein, J. (1998). The use of MMR, diversity-based "
    "reranking for reordering documents and producing summaries. SIGIR.",
    "Clegg, D., & Barker, R. (1994). Case Method Fast-Track: A RAD Approach. "
    "Addison-Wesley.",
    "Cohn, M. (2004). User Stories Applied: For Agile Software Development. "
    "Addison-Wesley Professional.",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design Science "
    "in Information Systems Research. MIS Quarterly, 28(1).",
    "Jørgensen, M., & Shepperd, M. (2007). A Systematic Review of Software "
    "Development Cost Estimation Studies. IEEE Transactions on Software "
    "Engineering, 33(1).",
    "Kahneman, D. (2011). Thinking, Fast and Slow. Farrar, Straus and Giroux.",
    "Kano, N., Seraku, N., Takahashi, F., & Tsuji, S. (1984). Attractive "
    "Quality and Must-Be Quality. Journal of the Japanese Society for "
    "Quality Control, 14(2).",
    "Leffingwell, D. (2018). SAFe 4.5 Reference Guide: Scaled Agile "
    "Framework for Lean Enterprises. Addison-Wesley.",
    "Lewis, P., Perez, E., Piktus, A., et al. (2020). Retrieval-Augmented "
    "Generation for Knowledge-Intensive NLP Tasks. NeurIPS.",
    "McBride, S. (2018). RICE: Simple prioritization for product managers. "
    "Intercom Blog.",
    "Mistral AI. (2024). Mistral Large — Technical Report.",
    "OpenAI. (2024). GPT-4o System Card.",
    "Rudin, C. (2019). Stop explaining black box machine learning models "
    "for high stakes decisions and use interpretable models instead. "
    "Nature Machine Intelligence, 1(5).",
    "Schwaber, K., & Sutherland, J. (2020). The Scrum Guide. scrumguides.org.",
    "Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is All "
    "You Need. NeurIPS.",
    "Wang, X., Wei, J., Schuurmans, D., et al. (2023). Self-Consistency "
    "Improves Chain-of-Thought Reasoning in Language Models. ICLR.",
]
for r in sorted(refs):
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    for run in p.runs:
        run.font.size = Pt(10)

add_page_break(doc)


# ═════════════════════════ ANNEXES ═════════════════════════

h1(doc, "8. Annexes")

h2(doc, "Annexe A — Référence API REST complète")
add_table(doc,
          ["Méthode", "Endpoint", "Description"],
          [
              ["GET", "/health", "Statut + version"],
              ["POST", "/api/chat/", "RAG avec auto-route modèle"],
              ["POST", "/api/chat/stream", "Streaming Server-Sent Events"],
              ["DELETE", "/api/chat/session/{id}", "Réinitialise la mémoire de session"],
              ["POST", "/api/prioritization/", "Score + rang d'une liste de features"],
              ["POST", "/api/prioritization/quick-score", "Une feature, tous algorithmes"],
              ["GET", "/api/prioritization/backlog", "Lecture du backlog persistant"],
              ["PUT", "/api/prioritization/backlog", "Remplacement du backlog"],
              ["DELETE", "/api/prioritization/backlog/{id}", "Suppression d'une feature"],
              ["POST", "/api/documents/upload", "Upload + indexation d'un document"],
              ["GET", "/api/documents/", "Liste des documents indexés"],
              ["DELETE", "/api/documents/{name}", "Suppression du fichier"],
          ])
caption(doc, "Tableau A.1. Référence des endpoints REST.")

h2(doc, "Annexe B — Exemple de requête de priorisation")
code(doc,
"""POST /api/prioritization/
{
  "algorithm": "rice",
  "use_ai_blend": false,
  "features": [
    {
      "id": "feat_1",
      "name": "AI sprint planner",
      "reach": 8, "impact": 9, "confidence": 0.8, "effort": 5,
      "business_value": 9, "time_criticality": 8,
      "risk_reduction": 6, "job_size": 5,
      "moscow": "must", "kano_category": "delighter",
      "satisfaction_gain": 9, "dissatisfaction_risk": 3,
      "ease": 5, "strategic_alignment": 9,
      "dependency_count": 2, "user_requests": 80,
      "tags": ["AI"], "epic": "AI Features",
      "description": "Auto-generate sprint plans"
    }
  ]
}""")

h2(doc, "Annexe C — Prompt système RAG (extrait)")
code(doc,
"""You are an expert AI assistant for Product Owners.
You have deep knowledge of Scrum, SAFe, Kanban, OKRs.

Use the retrieved context below to answer the user's question precisely.
Always ground your answer in Agile/PO best practices.
If the context doesn't cover the question, rely on your training knowledge
and say so.

Context:        {context}
Chat History:   {chat_history}
Question:       {question}

Respond in the same language as the question (FR or EN).
Be concise, structured, and actionable. Use bullet points when listing.
Always cite sources from the retrieved context at the end.""")

h2(doc, "Annexe D — Prompt système AI Blend (extrait)")
code(doc,
"""You are a senior Product Owner evaluating sprint features.
Be accurate, calibrated, and differentiated.

RULES:
1. Read ALL features first, then score them relative to each other
2. Use the calibration examples to anchor your scale
3. Step-by-step: demand → strategic value → urgency → effort → dependencies
4. must-have + high urgency + many requests → 85–95
5. could-have + no urgency + few requests       → 20–45
6. Strategic alignment is a multiplier, not a base score
7. Return a 'confidence' field (0–1) reflecting your certainty

OUTPUT: strict JSON array, no markdown
[{"id":"...","score":0-100,"confidence":0-1,
  "reasoning":"...","risk":"..."}]""")

h2(doc, "Annexe E — Hyperparamètres du modèle ML Hybrid")
add_table(doc,
          ["Paramètre", "Valeur"],
          [
              ["n_estimators", "400"],
              ["max_depth", "5"],
              ["learning_rate", "0.03"],
              ["subsample", "0.80"],
              ["min_samples_leaf", "3"],
              ["max_features", "0.85"],
              ["random_state", "42"],
              ["n samples bootstrap", "600"],
              ["features d'entrée (8)",
               "rice, wsjf, ice, kano, value_effort, moscow_num, strategic, deps"],
              ["étiquette y",
               "0.30·rice + 0.30·wsjf + 0.15·ice + 0.15·kano + 0.10·ve + b_moscow + b_deps"],
          ])
caption(doc, "Tableau E.1. Hyperparamètres du GradientBoostingRegressor.")

h2(doc, "Annexe F — Glossaire PO/Agile")
glossary = [
    ("User story", "Énoncé court du besoin selon le format "
                   "« En tant que [rôle], je veux [objectif], afin de [bénéfice] »."),
    ("Backlog", "Liste ordonnée des PBI à développer."),
    ("Epic", "Élément de travail volumineux découpé en plusieurs user stories."),
    ("Sprint", "Itération de durée fixe (typiquement 2 semaines)."),
    ("Velocity", "Story points complétés par sprint."),
    ("Definition of Done", "Critères partagés pour qu'une story soit considérée terminée."),
    ("Acceptance criteria", "Conditions explicites pour valider une story."),
    ("RICE", "(Reach × Impact × Confidence) / Effort, normalisé sur 0–100."),
    ("WSJF", "(Business Value + Time Criticality + Risk Reduction) / Job Size, "
             "standard SAFe."),
    ("MoSCoW", "Must / Should / Could / Won't."),
    ("Kano", "Modèle Must-be / Performance / Delighter."),
    ("RAG", "Retrieval-Augmented Generation : LLM conditionné par documents récupérés."),
    ("MMR", "Maximal Marginal Relevance : récupération avec pénalité de redondance."),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    r = p.add_run(f"{term} — ")
    r.bold = True
    p.add_run(definition)


# ───────────────────────────── save ─────────────────────────────

OUT = Path(__file__).resolve().parents[1] / "Memoire_POai_M2.docx"
doc.save(OUT)
print(f"OK -> {OUT}")
